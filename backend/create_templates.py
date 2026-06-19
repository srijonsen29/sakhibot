from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_fir_template():
    doc = Document()

    # title
    title = doc.add_heading("FIRST INFORMATION REPORT (FIR)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("To,")
    doc.add_paragraph("The Station House Officer,")
    doc.add_paragraph("{{police_station}} Police Station,")
    doc.add_paragraph("{{district}}, {{state}}")
    doc.add_paragraph("")

    doc.add_heading("Details of Complainant", level=2)
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = "Table Grid"
    fields1 = [
        ("Full Name", "{{complainant_name}}"),
        ("Address", "{{complainant_address}}"),
        ("Phone Number", "{{complainant_phone}}"),
        ("Age", "{{complainant_age}}"),
        ("Relationship to accused", "{{relationship}}"),
    ]
    for i, (label, value) in enumerate(fields1):
        table1.cell(i, 0).text = label
        table1.cell(i, 1).text = value

    doc.add_paragraph("")
    doc.add_heading("Details of Accused", level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    fields2 = [
        ("Name of Accused", "{{accused_name}}"),
        ("Address of Accused", "{{accused_address}}"),
        ("Relationship to complainant", "{{relationship}}"),
    ]
    for i, (label, value) in enumerate(fields2):
        table2.cell(i, 0).text = label
        table2.cell(i, 1).text = value

    doc.add_paragraph("")
    doc.add_heading("Incident Details", level=2)
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = "Table Grid"
    fields3 = [
        ("Date of Incident", "{{incident_date}}"),
        ("Place of Incident", "{{incident_place}}"),
        ("Nature of Incident", "{{incident_nature}}"),
    ]
    for i, (label, value) in enumerate(fields3):
        table3.cell(i, 0).text = label
        table3.cell(i, 1).text = value

    doc.add_paragraph("")
    doc.add_heading("Statement", level=2)
    doc.add_paragraph(
        "I, {{complainant_name}}, hereby state that on {{incident_date}}, "
        "at {{incident_place}}, {{accused_name}} committed the following act: "
        "{{incident_nature}}. I request you to kindly register this FIR and "
        "take necessary legal action under the relevant sections of IPC and "
        "Protection of Women from Domestic Violence Act, 2005."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Relevant Sections: IPC 498A, DV Act Section 12, CrPC Section 154")
    doc.add_paragraph("")
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph("")
    doc.add_paragraph("{{complainant_name}}")
    doc.add_paragraph("Date: {{incident_date}}")

    doc.save("templates/fir_template.docx")
    print("FIR template created.")


def create_dv_complaint_template():
    doc = Document()

    title = doc.add_heading(
        "APPLICATION UNDER SECTION 12 OF THE PROTECTION OF WOMEN "
        "FROM DOMESTIC VIOLENCE ACT, 2005", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("To,")
    doc.add_paragraph("The Hon'ble Judicial Magistrate / Metropolitan Magistrate,")
    doc.add_paragraph("{{court_district}}, {{state}}")
    doc.add_paragraph("")

    doc.add_heading("Aggrieved Person (Applicant)", level=2)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Full Name", "{{complainant_name}}"),
        ("Address", "{{complainant_address}}"),
        ("Phone", "{{complainant_phone}}"),
        ("Age", "{{complainant_age}}"),
    ]):
        table1.cell(i, 0).text = l
        table1.cell(i, 1).text = v

    doc.add_paragraph("")
    doc.add_heading("Respondent (Accused)", level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Full Name", "{{accused_name}}"),
        ("Address", "{{accused_address}}"),
        ("Relationship", "{{relationship}}"),
    ]):
        table2.cell(i, 0).text = l
        table2.cell(i, 1).text = v

    doc.add_paragraph("")
    doc.add_heading("Nature of Domestic Violence", level=2)
    doc.add_paragraph("{{incident_nature}}")

    doc.add_paragraph("")
    doc.add_heading("Relief Sought", level=2)
    doc.add_paragraph(
        "The applicant humbly prays for the following reliefs:\n"
        "1. Protection Order under Section 18 of DV Act\n"
        "2. Residence Order under Section 19 of DV Act\n"
        "3. Monetary Relief under Section 20 of DV Act\n"
        "4. Any other relief as deemed fit by this Hon'ble Court"
    )

    doc.add_paragraph("")
    doc.add_paragraph("Place: {{incident_place}}")
    doc.add_paragraph("Date: {{incident_date}}")
    doc.add_paragraph("")
    doc.add_paragraph("Applicant: {{complainant_name}}")

    doc.save("templates/dv_complaint.docx")
    print("DV complaint template created.")


def create_posh_complaint_template():
    doc = Document()

    title = doc.add_heading(
        "COMPLAINT UNDER THE SEXUAL HARASSMENT OF WOMEN AT "
        "WORKPLACE (PREVENTION, PROHIBITION AND REDRESSAL) ACT, 2013",
        level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("To,")
    doc.add_paragraph("The Presiding Officer,")
    doc.add_paragraph("Internal Complaints Committee (ICC),")
    doc.add_paragraph("{{organization_name}}")
    doc.add_paragraph("")

    doc.add_heading("Complainant Details", level=2)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Full Name", "{{complainant_name}}"),
        ("Designation", "{{designation}}"),
        ("Department", "{{department}}"),
        ("Phone / Email", "{{complainant_phone}}"),
    ]):
        table1.cell(i, 0).text = l
        table1.cell(i, 1).text = v

    doc.add_paragraph("")
    doc.add_heading("Respondent Details", level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Name of Accused", "{{accused_name}}"),
        ("Designation", "{{accused_designation}}"),
        ("Department", "{{accused_department}}"),
    ]):
        table2.cell(i, 0).text = l
        table2.cell(i, 1).text = v

    doc.add_paragraph("")
    doc.add_heading("Details of Incident", level=2)
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Date of Incident", "{{incident_date}}"),
        ("Place of Incident", "{{incident_place}}"),
    ]):
        table3.cell(i, 0).text = l
        table3.cell(i, 1).text = v

    doc.add_paragraph("")
    doc.add_heading("Description of Harassment", level=2)
    doc.add_paragraph("{{incident_nature}}")

    doc.add_paragraph("")
    doc.add_paragraph(
        "I request the ICC to conduct an inquiry into this matter as per "
        "Section 11 of the POSH Act 2013 and provide appropriate relief."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Date: {{incident_date}}")
    doc.add_paragraph("Signature: {{complainant_name}}")

    doc.save("templates/posh_complaint.docx")
    print("POSH complaint template created.")


if __name__ == "__main__":
    create_fir_template()
    create_dv_complaint_template()
    create_posh_complaint_template()
    print("\nAll 3 templates created in backend/templates/")